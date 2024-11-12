import streamlit as st
import time
import tempfile
from docxtpl import InlineImage
import pandas as pd
import calendar
from datetime import datetime
import os
import json
import matplotlib.pyplot as plt
from openpyxl.styles import Font
from collections import defaultdict
from docx import Document
from docx.shared import Pt, Inches, Mm
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
import pdfplumber
import re
import requests
from io import BytesIO
import tempfile
import zipfile
import shutil
from contextlib import contextmanager
from docxtpl import DocxTemplate, InlineImage
import logging
import platform
import subprocess
import markdown
from bs4 import BeautifulSoup
import folium
from streamlit_folium import st_folium
from geopy.geocoders import Nominatim
from streamlit_js_eval import get_geolocation
import google.generativeai as genai
from PIL import Image
import streamlit_lottie as st_lottie
from fuzzywuzzy import fuzz
from modulos.archivos import load_lottie_file, cargar_especies, cargar_ciudades
from modulos.donate import crear_donation_footer
# Importaciones espec√≠ficas de Windows
if platform.system() == 'Windows':
    import win32com.client as win32
    import pythoncom
else:
    # No importar pythoncom en sistemas no-Windows
    pass

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Configuraci√≥n de la p√°gina
st.set_page_config(
    page_title="Genera tu Entregable",
    page_icon="üé£",
    layout="wide"
)

# Diccionario para meses en espa√±ol
MESES_ES = {
    1: "Enero",
    2: "Febrero",
    3: "Marzo",
    4: "Abril",
    5: "Mayo",
    6: "Junio",
    7: "Julio",
    8: "Agosto",
    9: "Septiembre",
    10: "Octubre",
    11: "Noviembre",
    12: "Diciembre",
}

# Determinar la ruta base de la aplicaci√≥n
base_dir = os.path.dirname(os.path.abspath(__file__))
    
# Context manager para manejar la inicializaci√≥n y desinicializaci√≥n de COM
@contextmanager
def com_handler():
    pythoncom.CoInitialize()
    try:
        yield
    finally:
        pythoncom.CoUninitialize()

# Funci√≥n para obtener datos desde la API de SUNAT
def obtener_datos_sunat(dni):
    apisnet_key = st.secrets["APISNET"]["key"]
    url = f"https://api.apis.net.pe/v2/sunat/dni?numero={dni}&token={apisnet_key}"
    try:
        response = requests.get(url)
        if response.status_code == 200:
            data = response.json()
            nombres = f"{data.get('nombres', '')} {data.get('apellidoPaterno', '')} {data.get('apellidoMaterno', '')}".strip()
            # Generar iniciales
            iniciales = "".join(
                [nombre[0] for nombre in nombres.split() if nombre]
            ).upper()
            ruc = data.get("ruc", "")
            return nombres, iniciales, ruc
        else:
            st.error("Error al obtener datos de SUNAT. Verifica el DNI ingresado.")
            return None, None, None
    except Exception as e:
        st.error(f"Error al conectar con la API de SUNAT: {e}")
        return None, None, None

# Funci√≥n para generar el gr√°fico de barras
def generar_grafico(df_grafico, output_dir):
    plt.figure(figsize=(14, 8))  # Tama√±o ajustado para mejor visualizaci√≥n
    bar_width = 0.15
    days = df_grafico["D√≠a"]
    indices = range(len(days))
    num_months = len(df_grafico.columns) - 1  # Restando la columna 'D√≠a'
    colores = plt.cm.tab20.colors  # Usar una paleta de colores diferente

    # Crear barras para cada mes
    for i, mes in enumerate(df_grafico.columns[1:]):
        barras = plt.bar(
            [x + bar_width * i for x in indices],
            df_grafico[mes],
            width=bar_width,
            label=mes,
            color=colores[i % len(colores)],
        )
        # A√±adir etiquetas de datos
        for barra in barras:
            altura = barra.get_height()
            plt.text(
                barra.get_x() + barra.get_width() / 2.0,
                altura + 0.05,
                f"{altura}",
                ha="center",
                va="bottom",
                fontsize=8,
            )

    plt.xlabel("D√≠a", fontsize=14)
    plt.ylabel("Total KG", fontsize=14)
    plt.title("Descargas Diarias por Mes", fontsize=16)
    plt.xticks(
        [x + bar_width * (num_months - 1) / 2 for x in indices], days, rotation=45
    )
    plt.legend(title="Meses")
    plt.tight_layout()

    # Guardar el gr√°fico
    grafico_path = os.path.join(output_dir, "grafico_barras.png")
    plt.savefig(grafico_path, dpi=600)
    plt.close()
    return grafico_path

# Funci√≥n para convertir Word a PDF utilizando win32com.client con manejo de COM
def convertir_a_pdf(archivo_docx, archivo_pdf):
    try:
        logging.info(f"Convirtiendo {archivo_docx} a {archivo_pdf}")
        if not os.path.exists(archivo_docx):
            st.error(f"El archivo {archivo_docx} no existe y no se puede convertir a PDF.")
            logging.error(f"El archivo {archivo_docx} no existe.")
            return

        if platform.system() == 'Windows':
            with com_handler():
                word = win32.Dispatch('Word.Application')
                word.Visible = False
                word.DisplayAlerts = 0
                doc = word.Documents.Open(os.path.abspath(archivo_docx))
                time.sleep(2)
                doc.SaveAs(os.path.abspath(archivo_pdf), FileFormat=17)
                time.sleep(2)
                doc.Close()
                word.Quit()
        else:
            # Usar LibreOffice para la conversi√≥n en sistemas Linux
            try:
                # Intentar usar libreoffice
                subprocess.run([
                    'libreoffice', 
                    '--headless', 
                    '--convert-to', 
                    'pdf', 
                    archivo_docx,
                    '--outdir', 
                    os.path.dirname(archivo_pdf)
                ], check=True)
                
                # Renombrar el archivo si es necesario
                pdf_generado = os.path.splitext(archivo_docx)[0] + '.pdf'
                if pdf_generado != archivo_pdf:
                    os.rename(pdf_generado, archivo_pdf)
                    
            except subprocess.CalledProcessError:
                st.error("Error al convertir el documento. Aseg√∫rate de que LibreOffice est√© instalado.")
                return
            except Exception as e:
                st.error(f"Error inesperado durante la conversi√≥n: {str(e)}")
                return

        st.success(f"Documento convertido a PDF: {archivo_pdf}")
    except Exception as e:
        st.error(f"Error al convertir a PDF: {e}")
        logging.error(f"Error al convertir a PDF: {e}", exc_info=True)
        
# Funci√≥n para procesar y convertir documentos adicionales a PDF
def procesar_documentos_adicionales(archivos_docx, carpeta_salida, valores, base_dir):
    try:
        for docx_name in archivos_docx:
            # Ruta correcta a la plantilla
            archivo_entrada = os.path.join(base_dir, docx_name)
            archivo_salida_docx = docx_name  # Mismo nombre para salida
            archivo_salida_pdf = docx_name.replace(".docx", ".pdf")
            archivo_salida_docx_path = os.path.join(carpeta_salida, archivo_salida_docx)
            archivo_salida_pdf_path = os.path.join(carpeta_salida, archivo_salida_pdf)

            # Verificar si el archivo existe
            if not os.path.exists(archivo_entrada):
                st.warning(f"El archivo {archivo_entrada} no se encontr√≥.")
                continue

            # Llenar la plantilla de Word
            llenar_plantilla_word(
                archivo_entrada,
                archivo_salida_docx_path,
                valores,
                grafico_path=None,  # No se inserta gr√°fico aqu√≠
            )

            # Convertir a PDF
            convertir_a_pdf(archivo_salida_docx_path, archivo_salida_pdf_path)

        st.success(
            "Documentos adicionales procesados y convertidos a PDF correctamente."
        )
    except Exception as e:
        st.error(f"Error al procesar documentos adicionales: {e}")

# Funci√≥n para contar archivos en una lista
def contar_archivos(excel_files):
    try:
        return str(len(excel_files))
    except Exception as e:
        st.error(f"Error al contar los archivos: {e}")
        return "0"

# Funci√≥n para crear un archivo ZIP de los documentos generados
def crear_zip_archivos(archivos, carpeta_salida):
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "w") as zipf:
        for archivo in archivos:
            ruta_archivo = os.path.join(carpeta_salida, archivo)
            if os.path.exists(ruta_archivo):
                zipf.write(ruta_archivo, arcname=archivo)
    zip_buffer.seek(0)
    return zip_buffer

# Funci√≥n para generar recomendaciones por IA
rapidapi_key = st.secrets["RAPIDAPI"]["key"]

def generar_respuesta_rapidapi(prompt_text):
    url = "https://cheapest-gpt-4-turbo-gpt-4-vision-chatgpt-openai-ai-api.p.rapidapi.com/v1/chat/completions"

    payload = {
        "messages": [{"role": "user", "content": prompt_text}],
        "model": "gpt-4",
        "max_tokens": 1000,
        "temperature": 0.7,
    }
    headers = {
        "x-rapidapi-key": rapidapi_key,
        "x-rapidapi-host": "cheapest-gpt-4-turbo-gpt-4-vision-chatgpt-openai-ai-api.p.rapidapi.com",
        "Content-Type": "application/json",
    }

    try:
        response = requests.post(url, json=payload, headers=headers)
        response.raise_for_status()  # Raise an exception for bad status codes
        result = response.json()
        return result["choices"][0]["message"]["content"].strip()
    except requests.RequestException as e:
        st.error(f"Error al comunicarse con la API de RapidAPI: {e}")
        return ""
    except Exception as e:
        st.error(f"Error inesperado: {e}")
        return ""

# Funci√≥n para extraer el nombre del servicio desde el PDF
def extraer_nombre_servicio(pdf_file):
    texto_completo = ""
    try:
        with pdfplumber.open(BytesIO(pdf_file.read())) as pdf:
            for pagina in pdf.pages:
                texto_completo += pagina.extract_text()

        # Unificar el texto y eliminar saltos de l√≠nea excesivos
        texto_unido = ' '.join(texto_completo.split())

        # Patr√≥n para encontrar el texto entre "2. OBJETO DE LA CONTRATACION" y "3. FINALIDAD PUBLICA"
        patron = r'2\.\s*OBJETO\s*DE\s*LA\s*CONTRATACION\s*(.*?)\s*3\.\s*FINALIDAD\s*PUBLICA'

        # Buscar el patr√≥n en el texto unificado
        match = re.search(patron, texto_unido, re.DOTALL | re.IGNORECASE)

        if match:
            # Eliminar espacios adicionales y unificar el texto encontrado
            servicio = ' '.join(match.group(1).split())
            "Servicio extra√≠do:", servicio # Depuraci√≥n
            return servicio
        st.warning("Servicio no encontrado en el PDF.")
        return "Servicio no encontrado"
    except Exception as e:
        st.error(f"Error al extraer el servicio del PDF: {e}")
        st.exception(e)
        return "Servicio no encontrado"

# Funci√≥n para extraer las actividades desde el PDF
def extraer_actividades(pdf_file):
    texto_completo = ""
    try:
        # Reiniciar el buffer del archivo PDF
        pdf_file.seek(0)
        with pdfplumber.open(BytesIO(pdf_file.read())) as pdf:
            for pagina in pdf.pages:
                texto_completo += pagina.extract_text()

        # Unificar el texto y eliminar saltos de l√≠nea excesivos
        texto_unido = ' '.join(texto_completo.split())

        # Patr√≥n para encontrar el texto entre "7.1 ACTIVIDADES" y "7.2 PROCESO Y METODOLOG√çA"
        patron = r'7\.1\s*ACTIVIDADES\s*(.*?)\s*7\.2\s*PROCESO\s*Y\s*METODOLOG√çA'

        # Buscar el patr√≥n en el texto unificado
        match = re.search(patron, texto_unido, re.DOTALL | re.IGNORECASE)

        if match:
            # Extraer el texto encontrado entre las dos secciones
            actividades_texto = match.group(1)

            # Dividir el texto en base a los indicadores "a)", "b)", "c)", etc.
            actividades = re.split(r'\b([a-e])\)\s*', actividades_texto)

            # Crear variables para cada actividad
            act_a = act_b = act_c = act_d = act_e = None

            # Iterar sobre las actividades encontradas y asignarlas a las variables
            for i in range(1, len(actividades), 2):
                letra = actividades[i]
                contenido = ' '.join(actividades[i+1].split())
                if letra == 'a':
                    act_a = contenido
                elif letra == 'b':
                    act_b = contenido
                elif letra == 'c':
                    act_c = contenido
                elif letra == 'd':
                    act_d = contenido
                elif letra == 'e':
                    act_e = contenido

            st.write(f"Actividades extra√≠das: a): {act_a}, b): {act_b}, c): {act_c}, d): {act_d}, e): {act_e}")  # Depuraci√≥n
            return act_a, act_b, act_c, act_d, act_e

        st.warning("Actividades no encontradas en el PDF.")
        return (
            "Actividades no encontradas",
            "Actividades no encontradas",
            "Actividades no encontradas",
            "Actividades no encontradas",
            "Actividades no encontradas",
        )
    except Exception as e:
        st.error(f"Error al extraer las actividades del PDF: {e}")
        st.exception(e)
        return (
            "Actividades no encontradas",
            "Actividades no encontradas",
            "Actividades no encontradas",
            "Actividades no encontradas",
            "Actividades no encontradas",
        )

# Funci√≥n para reemplazar texto en p√°rrafos y tablas de manera robusta
def reemplazar_texto(doc, valores):
    """
    Reemplaza los marcadores de posici√≥n en todo el documento, incluyendo p√°rrafos y tablas.
    """
    # Reemplazo en p√°rrafos
    for paragraph in doc.paragraphs:
        for key, value in valores.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in paragraph.text:
                inline = paragraph.runs
                # Reemplazar en cada run
                for run in inline:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))
    
    # Reemplazo en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in valores.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in paragraph.text:
                            inline = paragraph.runs
                            for run in inline:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(value))

def limpiar_valor(valor):
    """
    Limpia y convierte el valor seg√∫n las siguientes reglas:
    - Si el valor es 'S/R' (ignorando may√∫sculas/min√∫sculas), devuelve 0.
    - Si el valor es una cadena que contiene ',' o ';', reemplaza estos caracteres por '.' y convierte a float.
    - Si el valor es num√©rico, lo devuelve como float.
    - En cualquier otro caso, devuelve 0.
    """
    if isinstance(valor, str):
        if valor.strip().upper() == "S/R":
            return 0.0
        # Reemplazar ',' y ';' por '.' y tratar de convertir a float
        valor = valor.replace(",", ".").replace(";", ".")
        try:
            return float(valor)
        except ValueError:
            return 0.0
    elif isinstance(valor, (int, float)):
        return float(valor)
    else:
        return 0.0

def generar_data_excel(excel_files, output_file):
    try:
        # Inicializar listas para almacenar datos de todas las hojas
        total_data = []
        comentarios_data = []
        grafico_data_list = []
        dfs_descripcion = []
        all_precios = pd.DataFrame()

        # Diccionario para los meses en espa√±ol
        meses_espanol = MESES_ES.copy()

        # Diccionario para almacenar los rangos de hora por fecha
        hora_range_data = {}

        # Bucle para procesar cada archivo Excel
        for file in excel_files:
            # Leer todas las hojas del archivo Excel
            xls = pd.ExcelFile(file, engine='openpyxl')
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet_name, header=None)

                # Procesamiento similar al actual pero para cada hoja
                # Aseg√∫rate de que cada hoja tenga la estructura esperada
                date = df.iloc[7, 11]  # Fecha en L8 (0-based)
                observation = df.iloc[12, 22]  # Observaci√≥n en W13 (0-based)
                total_value = df.iloc[11, 6]  # Total en G12 (0-based)

                # Obtener G13:G57 y X13:X57
                G_values = df.iloc[12:57, 6]  # G13:G57
                X_values = df.iloc[12:57, 23]  # X13:X57

                # Crear un DataFrame para Pescados y Mariscos
                total_items_df = pd.DataFrame({"G": G_values, "X": X_values})

                # Sumar los valores para Pescados y Mariscos
                sum_pescado = total_items_df.loc[
                    total_items_df["X"] == "PESCADOS", "G"
                ].sum()
                sum_marisco = total_items_df.loc[
                    total_items_df["X"] == "MARISCOS", "G"
                ].sum()

                # Manejar el valor total
                if pd.isna(total_value):
                    total_value = 0.0
                else:
                    try:
                        total_value = float(total_value)
                    except ValueError:
                        total_value = 0.0

                # Agregar los datos a la lista total_data
                total_data.append(
                    {
                        "Fecha": date,
                        "Pescado (KG)": sum_pescado,
                        "Marisco (KG)": sum_marisco,
                        "Total (KG)": total_value,
                    }
                )

                # Agregar datos para la hoja 'grafico'
                fecha_dt = (
                    pd.to_datetime(date, format="%d %m %Y", errors="coerce")
                    if pd.notna(date)
                    else pd.NaT
                )
                if pd.notna(fecha_dt):
                    dia = fecha_dt.day
                    mes_num = fecha_dt.month
                    mes_nombre = meses_espanol.get(mes_num, "Mes desconocido")
                    grafico_data_list.append(
                        {"D√≠a": dia, "Mes": mes_nombre, "Total_Kg": total_value}
                    )

                # Obtener 'Primer Valor' y '√öltimo Valor' de 'Hora'
                hora_col = df.iloc[12:57, 14]
                primer_valor = df.iloc[12, 14]
                primer_valor = "" if pd.isna(primer_valor) else primer_valor

                # Encontrar el √∫ltimo valor no nulo en la columna 'Hora'
                ultimo_valor = (
                    hora_col.dropna().iloc[-1]
                    if not hora_col.dropna().empty
                    else primer_valor
                )

                # Convertir 'primer_valor' y 'ultimo_valor' a formato 'HH:MM'
                def format_hora(valor):
                    if isinstance(valor, pd.Timestamp):
                        return valor.strftime("%H:%M")
                    elif isinstance(valor, str):
                        try:
                            return pd.to_datetime(valor).strftime("%H:%M")
                        except ValueError:
                            return valor
                    else:
                        return valor

                primer_valor_formateado = format_hora(primer_valor)
                ultimo_valor_formateado = format_hora(ultimo_valor)

                hora_range = f"{primer_valor_formateado} - {ultimo_valor_formateado}"

                # Almacenar el 'Hora' range por 'Fecha'
                hora_range_data[date] = hora_range

                # Leer columnas adicionales de 'Tama√±o' y 'Precio'
                tama√±o_columns = [15, 17, 19]  # P, R, T
                tama√±o_data = df.iloc[12:57, tama√±o_columns]
                
                # Aplicar limpieza a las columnas de 'Tama√±o'
                tama√±o_data_limpio = tama√±o_data.applymap(limpiar_valor)
                tama√±o_combined = tama√±o_data_limpio.stack().groupby(level=0).mean()

                # Convertir 'tama√±o_combined' a num√©rico y llenar NaN con 0
                tama√±o_combined = pd.to_numeric(tama√±o_combined, errors="coerce").fillna(0.0)

                precio_columns = [16, 18, 20]  # Q, S, U
                precio_data = df.iloc[12:57, precio_columns]
                
                # Aplicar limpieza a las columnas de 'Precio'
                precio_data_limpio = precio_data.applymap(limpiar_valor)
                
                # Incluir 'Nombre com√∫n' antes de melt para mantener la asociaci√≥n
                nombre_comun = df.iloc[
                    12:57, 2
                ].values  # Suponiendo que la columna 2 contiene 'Nombre com√∫n'
                precio_data_with_nombre = precio_data_limpio.copy()
                precio_data_with_nombre["Nombre com√∫n"] = nombre_comun

                # Realizar el melt incluyendo 'Nombre com√∫n' como id_vars
                precio_melted = precio_data_with_nombre.melt(
                    id_vars="Nombre com√∫n", value_vars=precio_columns, value_name="Precio"
                )

                # Convertir 'Precio' a num√©rico y filtrar valores > 0
                precio_melted["Precio"] = pd.to_numeric(
                    precio_melted["Precio"], errors="coerce"
                )
                precio_melted = precio_melted.dropna(subset=["Precio"])
                precio_melted = precio_melted[precio_melted["Precio"] > 0]

                # Agregar los precios individuales a 'all_precios'
                all_precios = pd.concat(
                    [all_precios, precio_melted[["Nombre com√∫n", "Precio"]]],
                    ignore_index=True,
                )

                # Crear DataFrame temporal para descripci√≥n
                temp_df = pd.DataFrame(
                    {
                        "Fecha": [date] * len(df.iloc[12:57, 2]),
                        "Nombre_com√∫n": df.iloc[12:57, 2],
                        "Nombre_cient√≠fico": df.iloc[12:57, 3],
                        "Volumen_Kg": df.iloc[12:57, 6],
                        "Procedencia": df.iloc[12:57, 8],
                        "Aparejo": df.iloc[12:57, 7],
                        "Hora": df.iloc[12:57, 14],
                        "Tama√±o": tama√±o_combined,
                        "Precio": precio_data_limpio.mean(axis=1),  # Promedio para descripci√≥n
                        "Observaci√≥n": "",
                    }
                )

                temp_df = temp_df.dropna(
                    subset=["Nombre_com√∫n", "Volumen_Kg", "Hora"], how="all"
                )

                if not temp_df.empty:
                    temp_df.iloc[0, temp_df.columns.get_loc("Observaci√≥n")] = observation
                    dfs_descripcion.append(temp_df)
                elif pd.notna(observation):
                    # Manejar casos donde solo hay observaci√≥n
                    observation_df = pd.DataFrame(
                        {
                            "Fecha": [date],
                            "Nombre_com√∫n": ["N/A"],
                            "Nombre_cient√≠fico": ["N/A"],
                            "Volumen_Kg": [0.0],
                            "Procedencia": ["N/A"],
                            "Aparejo": ["N/A"],
                            "Hora": ["N/A"],
                            "Tama√±o": [0.0],
                            "Precio": [0.0],
                            "Observaci√≥n": [observation],
                        }
                    )
                    dfs_descripcion.append(observation_df)
                    hora_range_data[date] = ""  # Sin rango de hora

        # Continuar con el procesamiento como ya lo tienes
        if not dfs_descripcion:
            st.error("No se encontraron datos de descripci√≥n para procesar.")
            return

        # Combinar todos los DataFrames de descripci√≥n
        result_df = pd.concat(dfs_descripcion, ignore_index=True)
        result_df["Fecha"] = pd.to_datetime(
            result_df["Fecha"], format="%d %m %Y", errors="coerce"
        )
        result_df = result_df.sort_values("Fecha")

        # Agrupar y resumir datos para descripci√≥n
        grouped_df = (
            result_df.groupby(["Fecha", "Nombre_com√∫n"])
            .agg(
                {
                    "Volumen_Kg": "sum",
                    "Tama√±o": "mean",
                    "Precio": "mean",  # Usamos el promedio para descripci√≥n
                    "Observaci√≥n": "first",
                    "Nombre_cient√≠fico": "first",
                }
            )
            .reset_index()
        )

        grouped_df["Tama√±o"] = grouped_df["Tama√±o"].round(2)
        grouped_df["Precio"] = grouped_df["Precio"].round(2)

        # Limpiar nombres de columnas para evitar espacios adicionales
        grouped_df.columns = grouped_df.columns.str.strip()

        # Formatear 'Fecha' para merge
        grouped_df["Fecha_str"] = grouped_df["Fecha"].dt.strftime("%d %m %Y")

        # Crear DataFrame de 'hora_range_data'
        hora_range_df = pd.DataFrame(
            list(hora_range_data.items()), columns=["Fecha_str", "Hora_Range"]
        )

        # Unir 'hora_range_df' con 'grouped_df'
        grouped_df = grouped_df.merge(hora_range_df, on="Fecha_str", how="left")
        grouped_df["Hora"] = grouped_df["Hora_Range"]
        grouped_df = grouped_df.drop(["Hora_Range", "Fecha_str"], axis=1)

        # Formatear 'Fecha' para visualizaci√≥n
        grouped_df["Fecha"] = grouped_df["Fecha"].dt.strftime("%d/%m/%Y")

        # Crear la columna 'Descripcion' antes de seleccionar las columnas
        def create_description(row):
            if (
                row["Nombre_com√∫n"] == "N/A"
                and row["Volumen_Kg"] == 0.0
                and row["Tama√±o"] == 0.0
                and row["Precio"] == 0.0
            ):
                return row["Observaci√≥n"]
            else:
                precio_formateado = "{:.2f}".format(row["Precio"])
                return f"{row['Nombre_com√∫n']} con {row['Volumen_Kg']} kg; talla promedio de {row['Tama√±o']} cm; precio por kilo de S/ {precio_formateado}"

        grouped_df["Descripcion"] = grouped_df.apply(create_description, axis=1)

        # Verificar que la columna 'Descripcion' se ha creado correctamente
        if "Descripcion" not in grouped_df.columns:
            raise KeyError("La columna 'Descripcion' no se ha creado correctamente.")

        # Definir el orden de las columnas seg√∫n lo solicitado
        columns_order = [
            "Fecha",
            "Nombre_com√∫n",
            "Volumen_Kg",
            "Hora",
            "Tama√±o",
            "Precio",
            "Observaci√≥n",
            "Descripcion",
        ]

        # Asegurarse de que todas las columnas existan antes de seleccionarlas
        missing_columns = [
            col for col in columns_order if col not in grouped_df.columns
        ]
        if missing_columns:
            raise KeyError(
                f"Las siguientes columnas faltan en 'grouped_df': {missing_columns}"
            )

        # Reordenar columnas
        grouped_df = grouped_df[columns_order]

        # Ordenar por 'Fecha' y 'Nombre_com√∫n' para asegurar el orden correcto
        grouped_df = grouped_df.sort_values(["Fecha", "Nombre_com√∫n"])

        # Establecer 'Hora' a cadena vac√≠a excepto para la primera ocurrencia de cada fecha
        grouped_df.loc[grouped_df.duplicated(subset="Fecha"), "Hora"] = ""

        # Crear 'total_df' con las nuevas columnas
        total_df = pd.DataFrame(total_data)
        total_df["Fecha"] = pd.to_datetime(
            total_df["Fecha"], format="%d %m %Y", errors="coerce"
        )
        total_df = total_df.sort_values("Fecha")

        # Convertir las columnas a num√©rico
        total_df["Pescado (KG)"] = pd.to_numeric(
            total_df["Pescado (KG)"], errors="coerce"
        ).fillna(0.0)
        total_df["Marisco (KG)"] = pd.to_numeric(
            total_df["Marisco (KG)"], errors="coerce"
        ).fillna(0.0)
        total_df["Total (KG)"] = pd.to_numeric(
            total_df["Total (KG)"], errors="coerce"
        ).fillna(0.0)

        # Formatear 'Fecha' para visualizaci√≥n
        total_df["Fecha"] = total_df["Fecha"].dt.strftime("%d/%m/%Y")

        # Calcular los totales
        total_sum_pescado = total_df["Pescado (KG)"].sum()
        total_sum_marisco = total_df["Marisco (KG)"].sum()
        total_sum_total = total_df["Total (KG)"].sum()

        # Crear la fila total
        total_row = pd.DataFrame(
            {
                "Fecha": ["TOTAL"],
                "Pescado (KG)": [total_sum_pescado],
                "Marisco (KG)": [total_sum_marisco],
                "Total (KG)": [total_sum_total],
            }
        )

        # Concatenar la fila total al DataFrame total_df
        total_df = pd.concat([total_df, total_row], ignore_index=True)

        # Reemplazar 0 con "‚Äì" en las columnas num√©ricas excepto en la fila 'TOTAL'
        numeric_columns = ["Pescado (KG)", "Marisco (KG)", "Total (KG)"]
        for col in numeric_columns:
            # Convertir a object para permitir strings
            total_df[col] = total_df[col].astype(object)
            # Crear m√°scara para filas que no sean 'TOTAL' y donde el valor sea 0
            mask = (total_df["Fecha"] != "TOTAL") & (total_df[col] == 0.0)
            # Asignar "‚Äì" donde la m√°scara es True
            total_df.loc[mask, col] = "‚Äì"

        # Filtrar registros v√°lidos (excluir N/A)
        valid_df = result_df[
            (result_df["Nombre_com√∫n"] != "N/A") & (result_df["Nombre_com√∫n"].notna())
        ]

        # Calcular el total de volumen por especie
        volumen_df = (
            valid_df.groupby("Nombre_com√∫n")
            .agg(KG_POR_ESPECIE=("Volumen_Kg", "sum"))
            .reset_index()
        )

        # Agrupar y calcular precios m√≠nimos y m√°ximos a partir de 'all_precios'
        if not all_precios.empty:
            # Filtrar precios > 0 para PRECIO MAS BAJO
            precios_mas_bajo_df = (
                all_precios[all_precios["Precio"] > 0]
                .groupby("Nombre com√∫n")
                .agg(PRECIO_MAS_BAJO=("Precio", "min"))
                .reset_index()
            )

            # Obtener PRECIO MAS ALTO
            precios_mas_alto_df = (
                all_precios.groupby("Nombre com√∫n")
                .agg(PRECIO_MAS_ALTO=("Precio", "max"))
                .reset_index()
            )

            # Combinar ambos DataFrames
            precios_df = pd.merge(
                precios_mas_bajo_df, precios_mas_alto_df, on="Nombre com√∫n", how="outer"
            )
        else:
            precios_df = pd.DataFrame(
                columns=["Nombre com√∫n", "PRECIO MAS BAJO (S/)", "PRECIO MAS ALTO (S/)"]
            )

        # Renombrar 'Nombre com√∫n' a 'Nombre_com√∫n' para consistencia
        precios_df.rename(columns={"Nombre com√∫n": "Nombre_com√∫n"}, inplace=True)

        # Unir todos los DataFrames
        especies_df = volumen_df.merge(precios_df, on="Nombre_com√∫n", how="left")

        # Agregar 'NOMBRE CIENTIFICO'
        nombres_cientificos = valid_df[
            ["Nombre_com√∫n", "Nombre_cient√≠fico"]
        ].drop_duplicates()
        especies_df = especies_df.merge(
            nombres_cientificos, on="Nombre_com√∫n", how="left"
        )

        # Renombrar columnas correctamente
        especies_df.rename(
            columns={
                "Nombre_com√∫n": "NOMBRE COMUN",
                "Nombre_cient√≠fico": "NOMBRE CIENTIFICO",
                "PRECIO_MAS_BAJO": "PRECIO MAS BAJO (S/)",
                "PRECIO_MAS_ALTO": "PRECIO MAS ALTO (S/)",
                "KG_POR_ESPECIE": "KG POR ESPECIE",
            },
            inplace=True,
        )

        # Limpiar nombres de columnas para evitar espacios adicionales
        especies_df.columns = especies_df.columns.str.strip()

        # Manejar valores NaN en precios
        if "PRECIO MAS BAJO (S/)" in especies_df.columns:
            especies_df["PRECIO MAS BAJO (S/)"] = especies_df[
                "PRECIO MAS BAJO (S/)"
            ].fillna("‚Äì")
        else:
            especies_df["PRECIO MAS BAJO (S/)"] = "‚Äì"

        if "PRECIO MAS ALTO (S/)" in especies_df.columns:
            especies_df["PRECIO MAS ALTO (S/)"] = especies_df[
                "PRECIO MAS ALTO (S/)"
            ].fillna("‚Äì")
        else:
            especies_df["PRECIO MAS ALTO (S/)"] = "‚Äì"

        # Definir el orden de las columnas seg√∫n lo solicitado
        especies_columns_order = [
            "NOMBRE COMUN",
            "NOMBRE CIENTIFICO",
            "PRECIO MAS BAJO (S/)",
            "PRECIO MAS ALTO (S/)",
            "KG POR ESPECIE",
        ]

        # Reordenar las columnas
        especies_df = especies_df[especies_columns_order]

        # Ordenar el DataFrame
        especies_df = especies_df.sort_values("NOMBRE COMUN")

        # Crear 'procedencia_df' excluyendo 'N/A' y valores nulos
        procedencia_df = result_df.dropna(subset=["Procedencia"])
        procedencia_df = procedencia_df[procedencia_df["Procedencia"] != "N/A"]
        procedencia_df = (
            procedencia_df.groupby("Procedencia")
            .agg(Volumen_Kg=("Volumen_Kg", "sum"))
            .reset_index()
        )
        procedencia_df = procedencia_df.sort_values("Procedencia")

        # Crear 'aparejos_data' agrupando 'result_df' por 'Aparejo'
        aparejos_df = (
            result_df.groupby("Aparejo")
            .agg(Volumen_Kg=("Volumen_Kg", "sum"))
            .reset_index()
        )
        aparejos_df = aparejos_df.sort_values("Aparejo")

        # Convertir 'aparejos_df' a diccionario
        aparejos_data = dict(zip(aparejos_df["Aparejo"], aparejos_df["Volumen_Kg"]))

        # Crear 'comentarios_df'
        # Especie principal
        if not especies_df.empty:
            especie_principal = especies_df.loc[especies_df["KG POR ESPECIE"].idxmax()]
            comentarios_data.extend(
                [
                    {
                        "Tipo_comentario": "Especie_principal_1",
                        "Comentario": f"{especie_principal['NOMBRE COMUN']}",
                    },
                    {
                        "Tipo_comentario": "Especie_principal_2",
                        "Comentario": f"({especie_principal['NOMBRE CIENTIFICO']})",
                    },
                    {
                        "Tipo_comentario": "Especie_principal_3",
                        "Comentario": f"{especie_principal['KG POR ESPECIE']}",
                    },
                ]
            )
        else:
            comentarios_data.extend(
                [
                    {
                        "Tipo_comentario": "Especie_principal_1",
                        "Comentario": "No hay datos disponibles",
                    },
                    {"Tipo_comentario": "Especie_principal_2", "Comentario": ""},
                    {"Tipo_comentario": "Especie_principal_3", "Comentario": ""},
                ]
            )

        # Especie m√≠nima (excluyendo KG POR ESPECIE igual a cero)
        especies_df_nonzero = especies_df[especies_df["KG POR ESPECIE"] > 0]
        if not especies_df_nonzero.empty:
            especie_minima = especies_df_nonzero.loc[
                especies_df_nonzero["KG POR ESPECIE"].idxmin()
            ]
            comentarios_data.extend(
                [
                    {
                        "Tipo_comentario": "Especie_minima_1",
                        "Comentario": f"{especie_minima['NOMBRE COMUN']}",
                    },
                    {
                        "Tipo_comentario": "Especie_minima_2",
                        "Comentario": f"({especie_minima['NOMBRE CIENTIFICO']})",
                    },
                    {
                        "Tipo_comentario": "Especie_minima_3",
                        "Comentario": f"{especie_minima['KG POR ESPECIE']}",
                    },
                ]
            )
        else:
            comentarios_data.extend(
                [
                    {
                        "Tipo_comentario": "Especie_minima_1",
                        "Comentario": "No hay datos disponibles",
                    },
                    {"Tipo_comentario": "Especie_minima_2", "Comentario": ""},
                    {"Tipo_comentario": "Especie_minima_3", "Comentario": ""},
                ]
            )

        # Procesar datos para mayor y menor d√≠a
        total_df_numeric = total_df.iloc[:-1].copy()  # Excluir la √∫ltima fila 'TOTAL'
        total_df_numeric["Total (KG)"] = pd.to_numeric(
            total_df_numeric["Total (KG)"], errors="coerce"
        )

        if not total_df_numeric.empty:
            # Encontrar m√°ximo total
            max_total = total_df_numeric["Total (KG)"].max()
            # Encontrar todas las fechas donde 'Total (KG)' es igual a max_total
            max_total_dates_df = total_df_numeric[
                total_df_numeric["Total (KG)"] == max_total
            ]
            max_total_dates_list = pd.to_datetime(
                max_total_dates_df["Fecha"], format="%d/%m/%Y", errors="coerce"
            ).dt.date.tolist()
            # Procesar fechas para crear 'mayor_dia_fecha'
            dates_info = []
            for date in max_total_dates_list:
                dia = date.day
                mes_num = date.month
                mes_nombre = meses_espanol.get(mes_num, "Mes desconocido")
                dates_info.append(
                    {"dia": dia, "mes_num": mes_num, "mes_nombre": mes_nombre}
                )

            # Agrupar fechas por mes
            dates_by_month = defaultdict(list)
            for info in dates_info:
                dates_by_month[info["mes_num"]].append(info)

            # Construir las cadenas de fechas
            date_strings = []
            for mes_num, infos in dates_by_month.items():
                dias = [str(info["dia"]) for info in infos]
                mes_nombre = infos[0]["mes_nombre"]
                if len(dias) > 1:
                    dias_str = ", ".join(dias[:-1]) + " y " + dias[-1]
                else:
                    dias_str = dias[0]
                date_strings.append(f"{dias_str} de {mes_nombre}")

            mayor_dia_fecha = "; ".join(date_strings)

            comentarios_data.append(
                {
                    "Tipo_comentario": "Mayor_dia",
                    "Comentario": f"{mayor_dia_fecha} se present√≥ la mayor descarga con un total de {max_total} kg",
                }
            )

            # Para 'Menor_dia' (excluyendo ceros)
            menor_dia_df = total_df_numeric[total_df_numeric["Total (KG)"] > 0]
            if not menor_dia_df.empty:
                min_total = menor_dia_df["Total (KG)"].min()
                # Encontrar todas las fechas donde 'Total (KG)' es igual a min_total
                min_total_dates_df = menor_dia_df[
                    menor_dia_df["Total (KG)"] == min_total
                ]
                min_total_dates_list = pd.to_datetime(
                    min_total_dates_df["Fecha"], format="%d/%m/%Y", errors="coerce"
                ).dt.date.tolist()
                # Procesar fechas para crear 'menor_dia_fecha'
                dates_info = []
                for date in min_total_dates_list:
                    dia = date.day
                    mes_num = date.month
                    mes_nombre = meses_espanol.get(mes_num, "Mes desconocido")
                    dates_info.append(
                        {"dia": dia, "mes_num": mes_num, "mes_nombre": mes_nombre}
                    )

                # Agrupar fechas por mes
                dates_by_month = defaultdict(list)
                for info in dates_info:
                    dates_by_month[info["mes_num"]].append(info)

                # Construir las cadenas de fechas
                date_strings = []
                for mes_num, infos in dates_by_month.items():
                    dias = [str(info["dia"]) for info in infos]
                    mes_nombre = infos[0]["mes_nombre"]
                    if len(dias) > 1:
                        dias_str = ", ".join(dias[:-1]) + " y " + dias[-1]
                    else:
                        dias_str = dias[0]
                    date_strings.append(f"{dias_str} de {mes_nombre}")

                menor_dia_fecha = "; ".join(date_strings)

                comentarios_data.append(
                    {
                        "Tipo_comentario": "Menor_dia",
                        "Comentario": f"{menor_dia_fecha} con {min_total} kg",
                    }
                )
            else:
                comentarios_data.append(
                    {
                        "Tipo_comentario": "Menor_dia",
                        "Comentario": "No se encontr√≥ un d√≠a con valor mayor a 0 kg",
                    }
                )
        else:
            comentarios_data.extend(
                [
                    {
                        "Tipo_comentario": "Mayor_dia",
                        "Comentario": "No hay datos disponibles para calcular el mayor d√≠a",
                    },
                    {
                        "Tipo_comentario": "Menor_dia",
                        "Comentario": "No hay datos disponibles para calcular el menor d√≠a",
                    },
                ]
            )

        # Aparejos
        aparejos_sorted = sorted(
            aparejos_data.items(), key=lambda x: x[1], reverse=True
        )
        aparejos_comment = " y ".join(
            [
                f"{aparejo} con {valor} Kg"
                for aparejo, valor in aparejos_sorted
                if valor > 0
            ]
        )
        comentarios_data.append(
            {"Tipo_comentario": "Aparejos", "Comentario": aparejos_comment}
        )

        # Para los comentarios de precios usando especies_df
        if not especies_df.empty:
            # Asegurarse de que los precios sean de tipo float
            especies_df["PRECIO MAS BAJO (S/)"] = pd.to_numeric(
                especies_df["PRECIO MAS BAJO (S/)"], errors="coerce"
            )
            especies_df["PRECIO MAS ALTO (S/)"] = pd.to_numeric(
                especies_df["PRECIO MAS ALTO (S/)"], errors="coerce"
            )

            # Encontrar el precio m√°s alto entre todos los registros con precio > 0
            max_precio = especies_df["PRECIO MAS ALTO (S/)"].max()
            max_precio_registros = especies_df[
                especies_df["PRECIO MAS ALTO (S/)"] == max_precio
            ]

            for _, row in max_precio_registros.iterrows():
                comentarios_data.extend(
                    [
                        {
                            "Tipo_comentario": "Mayor_Precio_1",
                            "Comentario": f"{row['NOMBRE COMUN']}",
                        },
                        {
                            "Tipo_comentario": "Mayor_Precio_2",
                            "Comentario": f"({row['NOMBRE CIENTIFICO']})",
                        },
                        {
                            "Tipo_comentario": "Mayor_Precio_3",
                            "Comentario": f"{row['PRECIO MAS ALTO (S/)']}",
                        },
                    ]
                )

            # Encontrar el precio m√°s bajo entre todos los registros con precio > 0
            min_precio = especies_df["PRECIO MAS BAJO (S/)"].min()
            min_precio_registros = especies_df[
                especies_df["PRECIO MAS BAJO (S/)"] == min_precio
            ]

            for _, row in min_precio_registros.iterrows():
                comentarios_data.extend(
                    [
                        {
                            "Tipo_comentario": "Menor_Precio_1",
                            "Comentario": f"{row['NOMBRE COMUN']}",
                        },
                        {
                            "Tipo_comentario": "Menor_Precio_2",
                            "Comentario": f"({row['NOMBRE CIENTIFICO']})",
                        },
                        {
                            "Tipo_comentario": "Menor_Precio_3",
                            "Comentario": f"{row['PRECIO MAS BAJO (S/)']}",
                        },
                    ]
                )
        else:
            comentarios_data.extend(
                [
                    {
                        "Tipo_comentario": "Mayor_Precio_1",
                        "Comentario": "No hay datos disponibles",
                    },
                    {"Tipo_comentario": "Mayor_Precio_2", "Comentario": ""},
                    {"Tipo_comentario": "Mayor_Precio_3", "Comentario": ""},
                    {
                        "Tipo_comentario": "Menor_Precio_1",
                        "Comentario": "No hay datos disponibles",
                    },
                    {"Tipo_comentario": "Menor_Precio_2", "Comentario": ""},
                    {"Tipo_comentario": "Menor_Precio_3", "Comentario": ""},
                ]
            )

        # Total de G12
        total_sum_G12 = sum(
            [
                d["Total (KG)"] if isinstance(d["Total (KG)"], (int, float)) else 0
                for d in total_data
            ]
        )
        comentarios_data.append(
            {"Tipo_comentario": "Total", "Comentario": f"{total_sum_G12} kg"}
        )

        # Crear DataFrame de comentarios
        comentarios_df = pd.DataFrame(comentarios_data)

        # Preparar datos para la hoja 'grafico'
        grafico_df = pd.DataFrame(grafico_data_list)

        # Crear pivot table sin reindexar a todos los d√≠as
        grafico_pivot = grafico_df.pivot_table(
            index="D√≠a", columns="Mes", values="Total_Kg", aggfunc="sum", fill_value=0
        ).reset_index()

        # Renombrar columnas de meses
        grafico_pivot.columns = ["D√≠a"] + [
            meses_espanol.get(mes_num, f"{mes_num}")
            for mes_num in grafico_pivot.columns[1:]
        ]

        # Ordenar por 'D√≠a'
        grafico_pivot = grafico_pivot.sort_values("D√≠a")

        # Convertir todos los valores a int para evitar problemas de formato en Excel
        grafico_pivot.iloc[:, 1:] = grafico_pivot.iloc[:, 1:].astype(int)

        # Guardar los DataFrames en el archivo Excel
        with pd.ExcelWriter(output_file, engine="openpyxl") as writer:
            # Hoja 'descripcion'
            grouped_df.to_excel(writer, index=False, sheet_name="descripcion")

            # Hoja 'total'
            total_df.to_excel(writer, index=False, sheet_name="total")

            # Hoja 'especies'
            especies_df.to_excel(writer, index=False, sheet_name="especies")

            # Hoja 'procedencia'
            procedencia_df.to_excel(writer, index=False, sheet_name="procedencia")

            # Hoja 'comentarios'
            comentarios_df.to_excel(writer, index=False, sheet_name="comentarios")

            # Hoja 'grafico'
            grafico_pivot.to_excel(writer, index=False, sheet_name="grafico")

            # Obtener el libro de trabajo
            workbook = writer.book

            # Formato para la hoja 'especies'
            if "especies" in workbook.sheetnames:
                worksheet = workbook["especies"]
                for row in worksheet.iter_rows(
                    min_row=2, min_col=2, max_col=2, max_row=worksheet.max_row
                ):
                    for cell in row:
                        cell.font = Font(italic=True)
                for col in ["A", "C", "D", "E"]:
                    for cell in worksheet[col + "1:" + col + str(worksheet.max_row)]:
                        cell[0].font = Font(bold=True)

            # Formato para la columna 'Fecha' en 'total'
            if "total" in workbook.sheetnames:
                worksheet = workbook["total"]
                for row in worksheet.iter_rows(
                    min_row=2, min_col=1, max_col=1, max_row=worksheet.max_row
                ):
                    for cell in row:
                        if cell.value != "TOTAL":
                            cell.number_format = "DD/MM/YYYY"

            # Formato para la hoja 'comentarios'
            if "comentarios" in workbook.sheetnames:
                worksheet = workbook["comentarios"]
                italic_font = Font(italic=True)
                for row in worksheet.iter_rows(
                    min_row=2, max_row=worksheet.max_row, min_col=1, max_col=2
                ):
                    if row[0].value in [
                        "Especie_principal_2",
                        "Especie_minima_2",
                        "Mayor_Precio_2",
                        "Menor_Precio_2",
                    ]:
                        row[1].font = italic_font

            # Formato para la hoja 'grafico'
            if "grafico" in workbook.sheetnames:
                worksheet = workbook["grafico"]
                # Ajustar el ancho de las columnas para mejor legibilidad
                for column_cells in worksheet.columns:
                    length = max(len(str(cell.value)) for cell in column_cells)
                    adjusted_width = length + 2
                    worksheet.column_dimensions[column_cells[0].column_letter].width = (
                        adjusted_width
                    )

        st.success(f"Archivo generado: {output_file}")
        return True
    except Exception as e:
        st.error(f"Error al analizar los datos: {e}")
        return False

def llenar_primera_tabla(doc, df_descripcion):
    if df_descripcion is None or df_descripcion.empty:
        logging.error("El DataFrame 'df_descripcion' es None o est√° vac√≠o.")
        return 

    try:
        table = doc.tables[0]  # Primera tabla 

        # Limpiar filas existentes excepto el encabezado 
        while len(table.rows) > 1:
            table._element.remove(table.rows[-1]._element)

        # Agrupar por fecha y concatenar descripciones 
        df_grouped = df_descripcion.groupby('Fecha').agg({ 
            'Hora': 'first',  # Tomar la primera hora de cada fecha 
            'Descripcion': lambda x: '\n'.join(x)  # Concatenar descripciones 
        }).reset_index() 

        # Convertir fecha al formato correcto 
        df_grouped['Fecha'] = pd.to_datetime(df_grouped['Fecha'], format='%d/%m/%Y').dt.strftime('%d/%m/%Y') 

        # Llenar la tabla 
        for _, row in df_grouped.iterrows(): 
            new_row = table.add_row().cells 
            new_row[0].text = row['Fecha'] 
            new_row[1].text = str(row['Hora']) if pd.notnull(row['Hora']) else '' 
            new_row[2].text = str(row['Descripcion']) 

            # Aplicar formato 
            for cell in new_row: 
                paragraph = cell.paragraphs[0] 
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT 
                for run in paragraph.runs: 
                    run.font.size = Pt(10) 

    except Exception as e:
        logging.error(f"Error al llenar la primera tabla: {e}")
        st.error(f"Error al llenar la primera tabla: {e}")

def llenar_segunda_tabla(doc, df_total):
    if df_total is None or df_total.empty:
        logging.error("Linea 1323: El DataFrame 'df_total' es None o est√° vac√≠o.")
        return

    try:
        table = doc.tables[1]  # Segunda tabla

        # Limpiar filas existentes excepto el encabezado
        for row in table.rows[1:]:
            tbl = table._tbl
            tbl.remove(row._tr)

        # Llenar la tabla
        for _, row in df_total.iterrows():
            new_row = table.add_row().cells
            new_row[0].text = str(row['Fecha'])
            new_row[1].text = str(row['Pescado (KG)'])
            new_row[2].text = str(row['Marisco (KG)'])
            new_row[3].text = str(row['Total (KG)'])

            # Centrar valores
            for cell in new_row:
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    except Exception as e:
        logging.error(f"Error al llenar la segunda tabla: {e}")

def llenar_tercera_tabla(doc, df_especies):
    try:
        table = doc.tables[2]  # Tercera tabla
        
        # Limpiar filas existentes excepto el encabezado
        for row in table.rows[1:]:
            tbl = table._tbl
            tbl.remove(row._tr)

        # Seleccionar solo las columnas necesarias
        columnas = ['NOMBRE COMUN', 'NOMBRE CIENTIFICO', 'PRECIO MAS BAJO (S/)', 'PRECIO MAS ALTO (S/)']
        df_especies = df_especies[columnas]

        # Llenar la tabla
        for _, row in df_especies.iterrows():
            new_row = table.add_row().cells
            for i, value in enumerate(row):
                new_row[i].text = str(value)
                
                # Aplicar formato
                paragraph = new_row[i].paragraphs[0]
                if i < 2:  # Nombre com√∫n y cient√≠fico alineados a la izquierda
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:  # Precios centrados
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    except Exception as e:
        logging.error(f"Error al llenar la tercera tabla: {e}")

def llenar_tablas_procedencia(doc, df_procedencia):
    try:
        # Llenar cuarta y quinta tabla (sin encabezado)
        for table_index in [3, 4]:
            table = doc.tables[table_index]
            
            # Limpiar todas las filas (no hay encabezado)
            for row in table.rows[:]:
                tbl = table._tbl
                tbl.remove(row._tr)

            # Llenar la tabla desde la primera fila
            for _, row in df_procedencia.iterrows():
                new_row = table.add_row().cells
                new_row[0].text = str(row['Procedencia'])
                new_row[1].text = str(row['Volumen_Kg'])

                # Alinear texto
                new_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                new_row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    except Exception as e:
        logging.error(f"Error al llenar las tablas de procedencia: {e}")

def add_species_images_and_captions(doc, uploaded_images, imagenes_dir, especies_df):
    if not uploaded_images:
        logging.warning("No hay im√°genes para insertar.")
        return

    doc.add_paragraph("FOTOGRAF√çAS").style = 'Heading 1'
    fig_num = 2  # Iniciar numeraci√≥n de figuras 

    for i, image in enumerate(uploaded_images):
        try:
            imagen_path = os.path.join(imagenes_dir, image.name)

            # Verificar si la imagen existe y es v√°lida
            if not os.path.exists(imagen_path):
                st.error(f"La imagen {image.name} no se encontr√≥ en {imagenes_dir}.")
                logging.error(f"La imagen {image.name} no se encontr√≥ en {imagenes_dir}.")
                continue

            # Intentar abrir la imagen para verificar su integridad
            try:
                with Image.open(imagen_path) as img:
                    img.verify()  # Verificar la integridad de la imagen
            except Exception as e:
                st.error(f"Error al verificar la imagen {image.name}: {e}")
                logging.error(f"Error al verificar la imagen {image.name}: {e}")
                continue

            # Obtener la especie seleccionada por el usuario del session_state 
            especie_key = f"especie_{i}"
            if especie_key in st.session_state:
                # La selecci√≥n est√° en formato "NOMBRE COMUN (NOMBRE CIENTIFICO)" 
                seleccion = st.session_state[especie_key] 
                nombre_comun = seleccion.split(" (")[0] 
                nombre_cientifico = seleccion.split("(")[1].rstrip(")") 

                # Buscar los datos correspondientes en especies_df 
                especie_data = especies_df[especies_df['NOMBRE COMUN'] == nombre_comun] 

                if not especie_data.empty: 
                    kg_especie = especie_data['KG POR ESPECIE'].iloc[0] 
                else: 
                    kg_especie = 0 
            else: 
                nombre_comun = "Especie no identificada" 
                nombre_cientifico = "Nombre cient√≠fico no disponible" 
                kg_especie = 0 

            # A√±adir la imagen 
            p = doc.add_paragraph() 
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER 
            run = p.add_run() 
            run.add_picture(imagen_path, width=Inches(4.331)) 

            # A√±adir el pie de foto con el formato especificado 
            caption = doc.add_paragraph() 
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER 

            # Primera parte del pie de foto (Fig. XX y nombre com√∫n) 
            caption_run = caption.add_run(f"Fig. {fig_num:02d} Ejemplares de {nombre_comun} (") 
            caption_run.italic = False 

            # Nombre cient√≠fico en cursiva 
            scientific_name_run = caption.add_run(f"{nombre_cientifico}") 
            scientific_name_run.italic = True 

            # Parte final con los kg 
            kg_run = caption.add_run(f") tuvo una descarga de {kg_especie} Kg") 
            kg_run.italic = False 

            fig_num += 1 

        except Exception as e:
            st.error(f"Error procesando la imagen {image.name}: {e}")
            logging.error(f"Error procesando la imagen {image.name}: {e}")

def obtener_comentarios_excel(comentarios_df):
    comentarios = {}
    try:
        for _, row in comentarios_df.iterrows():
            tipo = row['Tipo_comentario']
            comentario = row['Comentario']
            comentarios[tipo] = comentario
        return comentarios
    except Exception as e:
        st.error(f"Error al obtener comentarios: {e}")
        return {}
    
def llenar_plantilla_word(archivo_entrada, archivo_salida, valores, grafico_path=None, df_descripcion=None, df_total=None, df_especies=None, df_procedencia=None, uploaded_images=None, imagenes_dir=None): 
    try: 
        # Cargar la plantilla de Word con docxtpl 
        doc = DocxTemplate(archivo_entrada) 

        # Manejar la firma 
        # if valores.get('firma'): 
        #     try: 
        #         # Verificar si 'firma' es una ruta de archivo existente 
        #         if isinstance(valores['firma'], str) and os.path.exists(valores['firma']): 
        #             # Redimensionar la imagen manteniendo la relaci√≥n de aspecto
        #             imagen = Image.open(valores['firma'])
        #             # Calcular el nuevo ancho manteniendo la proporci√≥n
        #             ancho_maximo = 2 * 72  # Aproximadamente 2 pulgadas
        #             proporcion = ancho_maximo / imagen.width
        #             nuevo_alto = int(imagen.height * proporcion)
                    
        #             # Insertar la imagen de firma con dimensiones espec√≠ficas
        #             valores['firma'] = InlineImage(doc, valores['firma'], width=Inches(2), height=Inches(nuevo_alto/72))
        #         elif isinstance(valores['firma'], InlineImage): 
        #             # Si ya es un InlineImage, dejarlo como est√° 
        #             pass 
        #         else: 
        #             st.warning("La ruta de la firma no es v√°lida.") 
        #             valores['firma'] = None 
        #     except Exception as e: 
        #         st.error(f"Error al procesar la imagen de firma: {e}") 
        #         valores['firma'] = None 

        # Si hay un gr√°fico para insertar, agr√©galo al contexto
        if grafico_path:
            valores['grafico'] = InlineImage(doc, grafico_path, width=Mm(130))

        # Renderizar la plantilla
        doc.render(valores)
        
        # Guardar el documento generado
        doc.save(archivo_salida)
        st.success(f"Documento generado: {archivo_salida}")
        
        # Abrir el documento con python-docx para aplicar formatos adicionales
        doc_py = Document(archivo_salida)
        
        # Aplicar formato a las recomendaciones
        for paragraph in doc_py.paragraphs:
            if '{{recomendaciones}}' in paragraph.text:
                # Limpiar el p√°rrafo existente
                paragraph.clear()
                
                # Procesar el texto markdown
                markdown_text = valores.get('recomendaciones', '')
                html = markdown.markdown(markdown_text)
                soup = BeautifulSoup(html, 'html.parser')
                
                for element in soup.descendants:
                    if element.name == 'p':
                        # Nuevo p√°rrafo para cada <p>
                        if paragraph.text:
                            paragraph = doc_py.add_paragraph()
                        for child in element.children:
                            run = paragraph.add_run(child.string if child.string else '')
                            if child.name == 'strong':
                                run.bold = True
                            elif child.name == 'em':
                                run.italic = True
                    elif element.name == 'ul':
                        for li in element.find_all('li'):
                            bullet_para = doc_py.add_paragraph(li.text, style='List Bullet')

        # Continuar con el resto del procesamiento...
        if df_descripcion is not None:
            llenar_primera_tabla(doc_py, df_descripcion)
        if df_total is not None:
            llenar_segunda_tabla(doc_py, df_total)
        if df_especies is not None:
            llenar_tercera_tabla(doc_py, df_especies)
        if df_procedencia is not None:
            llenar_tablas_procedencia(doc_py, df_procedencia)
        if uploaded_images is not None and imagenes_dir is not None:
            add_species_images_and_captions(doc_py, uploaded_images, imagenes_dir, df_especies)

        # Guardar los cambios finales
        doc_py.save(archivo_salida)
        st.success("Documento formateado y guardado correctamente.")
        logging.info("Documento formateado y guardado correctamente.")

    except Exception as e:
        st.error(f"Error al llenar la plantilla de Word: {e}")
        logging.error(f"Error al llenar la plantilla de Word: {e}", exc_info=True)

def leer_y_formatear_dataframes(excel_bytes_io):
    try:
        # Verificar que excel_bytes_io tiene datos
        excel_bytes_io.seek(0, os.SEEK_END)
        size = excel_bytes_io.tell()
        if size == 0:
            st.error("El archivo data.xlsx est√° vac√≠o.")
            return ""
        excel_bytes_io.seek(0)  # Reiniciar el cursor
        
        logging.info(f"Tipo de excel_bytes_io: {type(excel_bytes_io)}")
        excel_file = pd.ExcelFile(excel_bytes_io, engine='openpyxl')
        logging.info("ExcelFile creado correctamente")

        # Leer cada hoja espec√≠fica
        descripcion_df = excel_file.parse(sheet_name="descripcion")
        total_df = excel_file.parse(sheet_name="total")
        especies_df = excel_file.parse(sheet_name="especies")
        procedencia_df = excel_file.parse(sheet_name="procedencia")
        comentarios_df = excel_file.parse(sheet_name="comentarios")
        grafico_df = excel_file.parse(sheet_name="grafico")
        
        # Convertir cada DataFrame a texto
        descripcion_text = descripcion_df.to_string(index=False)
        total_text = total_df.to_string(index=False)
        especies_text = especies_df.to_string(index=False)
        procedencia_text = procedencia_df.to_string(index=False)
        comentarios_text = comentarios_df.to_string(index=False)
        grafico_text = grafico_df.to_string(index=False)
        
        # Combinar todos los textos en una sola cadena
        dataframes_text = f"""
        Descripci√≥n:
        {descripcion_text}

        Totales:
        {total_text}

        Especies:
        {especies_text}

        Procedencia:
        {procedencia_text}

        Comentarios:
        {comentarios_text}

        Gr√°fico:
        {grafico_text}
        """
        return dataframes_text
    except Exception as e:
        st.error(f"Error al leer y formatear data.xlsx: {e}")
        logging.error(f"Error al leer y formatear data.xlsx: {e}")
        return ""

def markdown_to_docx_text(markdown_text):
    html = markdown.markdown(markdown_text)
    soup = BeautifulSoup(html, 'html.parser')
    text = soup.get_text()
    text = text.replace('‚Ä¢', '¬∑')
    text = text.replace('\n\n', '\n')
    return text

def obtener_direccion_desde_coordenadas(lat, lon):
    geolocator = Nominatim(user_agent="my_streamlit_app")
    try:
        location = geolocator.reverse((lat, lon))
        return location.address if location else None
    except Exception as e:
        st.error(f"Error al obtener la direcci√≥n: {e}")
        return None

def crear_mapa(lat=None, lon=None, zoom=13):
    # Usar coordenadas proporcionadas o predeterminadas de Lima, Per√∫
    if lat is None or lon is None:
        lat, lon = -12.0464, -77.0428  # Coordenadas de Lima

    # Crear el mapa centrado en la ubicaci√≥n y usando el zoom proporcionado
    m = folium.Map(location=[lat, lon], zoom_start=zoom)

    # Agregar marcador si hay coordenadas espec√≠ficas
    folium.Marker(
        [lat, lon],
        popup="Ubicaci√≥n actual",
        icon=folium.Icon(color='red', icon='info-sign'),
        draggable=False
    ).add_to(m)

    return m

def seccion_direccion():
    # Inicializar variables de sesi√≥n si no existen
    if 'lat' not in st.session_state:
        st.session_state.lat = -12.0464  # Lima por defecto
    if 'lon' not in st.session_state:
        st.session_state.lon = -77.0428
    if 'zoom' not in st.session_state:
        st.session_state.zoom = 13
    if 'direccion' not in st.session_state:
        st.session_state.direccion = ''

    # Columnas para organizar la secci√≥n
    col1, col2 = st.columns([2, 3])

    with col1:
        # Recuperar la ubicaci√≥n desde st.session_state
        if 'geo_loc' in st.session_state:
            loc = st.session_state['geo_loc']
            if loc and 'coords' in loc:
                st.session_state.lat = loc['coords']['latitude']
                st.session_state.lon = loc['coords']['longitude']
                direccion = obtener_direccion_desde_coordenadas(
                    st.session_state.lat,
                    st.session_state.lon
                )
                if direccion:
                    st.session_state.direccion = direccion

        # Un solo campo de direcci√≥n fuera de las columnas
        direccion_input = st.text_input(
            "Direcci√≥n",
            value=st.session_state.get('direccion', ''),
            key="direccion_input",
            max_chars=100
        )
        if st.button("Obtener ubicaci√≥n actual"):
            get_geolocation('geo_loc')
            st.write("Obteniendo ubicaci√≥n...")
        # Actualizar el estado con el valor del input
        st.session_state.direccion = direccion_input

    with col2:
        # Mostrar el mapa con la ubicaci√≥n si est√° disponible y el zoom actual
        mapa = crear_mapa(
            lat=st.session_state['lat'],
            lon=st.session_state['lon'],
            zoom=st.session_state['zoom']
        )
        mapa_data = st_folium(
            mapa,
            height=300,
            width=None,
            returned_objects=["last_clicked", "zoom"]
        )

        # Actualizar ubicaci√≥n cuando se hace clic en el mapa
        if mapa_data["last_clicked"]:
            clicked_lat = mapa_data["last_clicked"]["lat"]
            clicked_lng = mapa_data["last_clicked"]["lng"]

            # Guardar el zoom actual antes de actualizar
            if mapa_data.get("zoom"):
                st.session_state['zoom'] = mapa_data["zoom"]

            # Actualizar estado
            st.session_state['lat'] = clicked_lat
            st.session_state['lon'] = clicked_lng
            nueva_direccion = obtener_direccion_desde_coordenadas(clicked_lat, clicked_lng)
            if nueva_direccion:
                st.rerun()
        # Actualizar el zoom incluso si no se hace clic
        elif mapa_data.get("zoom"):
            st.session_state['zoom'] = mapa_data["zoom"]

# def seccion_firma(): 
#     upload_key = "firma_upload_entregable" 
#     preview_key = "firma_preview_entregable" 

#     if 'firma_file' not in st.session_state: 
#         st.session_state.firma_file = None 

#     firma_file = st.file_uploader( 
#         "Selecciona tu imagen de firma", 
#         type=["png", "jpg", "jpeg"], 
#         key=upload_key 
#     ) 

#     if firma_file is not None: 
#         with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(firma_file.name)[1]) as tmp_file: 
#             tmp_file.write(firma_file.getbuffer()) 
#             firma_path = tmp_file.name 

#         st.session_state.firma_file = firma_path 

#         # Muestra vista previa 
#         image = Image.open(firma_file) 
#         st.image(image, caption="Vista previa de la firma", width=300) 

#         return firma_path  # Devolver la ruta del archivo 

#     return st.session_state.firma_file 

def main():
    st.title("Genera tu Entregable")

    # Inicializar 'recomendaciones', 'data_excel', 'servicio' y actividades en session_state si no existen
    if 'recomendaciones' not in st.session_state:
        st.session_state['recomendaciones'] = ""
    if 'data_excel' not in st.session_state:
        st.session_state['data_excel'] = None
        logging.info(f"Tipo de st.session_state['data_excel']: {type(st.session_state['data_excel'])}")
    if 'servicio' not in st.session_state:
        st.session_state['servicio'] = ""
    if 'actividades' not in st.session_state:
        st.session_state['actividades'] = {"a": "", "b": "", "c": "", "d": "", "e": ""}
    if 'direccion' not in st.session_state:
        st.session_state['direccion'] = ''

    # Subida de archivos Excel
    st.header("Sube tus reportes Excel")
    excel_files = st.file_uploader(
        "Selecciona tus archivos Excel",
        type=["xlsx", "xls"],
        accept_multiple_files=True,
    )

    # Subida de PDF
    st.header("Sube tu TDR (PDF)")
    pdf_file = st.file_uploader("Selecciona tu archivo PDF", type=["pdf"])

# Subida de im√°genes con selecci√≥n de especie
    st.header("Adjuntar fotos de las especies")
    uploaded_images = st.file_uploader("Seleccionar fotos", type=["png", "jpg", "jpeg"], accept_multiple_files=True)

    especies = cargar_especies()

    if uploaded_images:
        tipos_especies = list(set(especie["por_tipo"] for especie in especies))
        
        for i, image in enumerate(uploaded_images):
            col1, col2 = st.columns([2, 3])
            with col1:
                st.image(image, caption=f"Imagen {i+1}", width=200)
            with col2:
                # Extraer el nombre base de la imagen sin extensi√≥n y en min√∫sculas
                base_name = os.path.splitext(image.name)[0].lower()
                
                # Limpiar el nombre eliminando sufijos como '- copia'
                base_name = re.sub(r'\s*-\s*copia$', '', base_name)
                
                # Separar m√∫ltiples nombres por coma
                nombres_archivo = [nombre.strip() for nombre in base_name.split(",")]

                # Inicializar variables
                tipo_seleccionado = "Desconocido"
                especie_seleccionada = "Especie no identificada"
                nombre_cientifico = "Nombre cient√≠fico no disponible"

                # Buscar coincidencia en especies.json
                for nombre in nombres_archivo:
                    especie_info = next(
                        (especie for especie in especies 
                         if nombre in [nombre_comun.strip().lower() for nombre_comun in especie["por_nombres_comunes"].split(",")]),
                        None
                    )
                    if especie_info:
                        tipo_seleccionado = especie_info["por_tipo"]
                        especie_seleccionada = f"{especie_info['por_nombres_comunes']} ({especie_info['por_nombre_cientifico']})"
                        nombre_cientifico = especie_info["por_nombre_cientifico"]
                        break  # Detenerse al encontrar la primera coincidencia

                if not especie_info:
                    # Si no se encuentra una coincidencia, usar IA para detectar la especie
                    st.write("Analizando la imagen para determinar la especie...")
                    try:
                        # Inicializar el modelo Gemini
                        genai.configure(api_key=st.secrets["GEMINIAPI"]["key"])
                        model = genai.GenerativeModel(model_name="gemini-1.5-flash")
                        
                        # Convertir la imagen a bytes
                        image_bytes = image.read()
                        image_pil = Image.open(BytesIO(image_bytes))
                        # Modificar el prompt para ser m√°s espec√≠fico
                        prompt = """Identifica la especie en la imagen. 
                        Si es un pez, indica 'PESCADO:' seguido del nombre.
                        Si es un marisco (crust√°ceo, molusco), indica 'MARISCO:' seguido del nombre.
                        Si es un alga u otra especie marina, indica 'ESPECIE:' seguido del nombre."""

                        # Generar la respuesta de Gemini AI
                        response = model.generate_content([prompt, image_pil], stream=True)
                        response.resolve()
                        respuesta_ai = response.text.upper().strip()

                        # Determinar el tipo basado en el prefijo de la respuesta
                        tipo_detectado = None
                        especie_detectada = None

                        if "PESCADO:" in respuesta_ai:
                            tipo_detectado = "PESCADOS"
                            especie_detectada = respuesta_ai.split("PESCADO:")[1].strip()
                        elif "MARISCO:" in respuesta_ai:
                            tipo_detectado = "MARISCOS"
                            especie_detectada = respuesta_ai.split("MARISCO:")[1].strip()
                        elif "ESPECIE:" in respuesta_ai:
                            tipo_detectado = "ESPECIES"
                            especie_detectada = respuesta_ai.split("ESPECIE:")[1].strip()
                        else:
                            # Si no hay prefijo, intentar determinar el tipo por palabras clave
                            especie_detectada = respuesta_ai
                            if any(palabra in respuesta_ai for palabra in ["CANGREJO", "LANGOSTA", "CAMAR√ìN", "ALMEJA", "OSTRA"]):
                                tipo_detectado = "MARISCOS"
                            elif any(palabra in respuesta_ai for palabra in ["ALGA", "CORAL", "PLANCTON"]):
                                tipo_detectado = "ESPECIES"
                            else:
                                tipo_detectado = "PESCADOS"  # Default solo si no hay otras coincidencias

                        # Buscar la especie m√°s similar en el JSON basado en el tipo detectado
                        especies_del_tipo = [especie for especie in especies if especie["por_tipo"] == tipo_detectado]
                        
                        mejor_coincidencia = None
                        mejor_ratio = 0
                        
                        for especie in especies_del_tipo:
                            nombres_comunes = [nombre.strip().upper() for nombre in especie["por_nombres_comunes"].split(",")]
                            for nombre in nombres_comunes:
                                ratio = fuzz.ratio(especie_detectada, nombre)
                                if ratio > mejor_ratio and ratio > 60:  # Umbral de similitud del 60%
                                    mejor_ratio = ratio
                                    mejor_coincidencia = especie

                        if mejor_coincidencia:
                            tipo_seleccionado = mejor_coincidencia["por_tipo"]
                            especie_seleccionada = f"{mejor_coincidencia['por_nombres_comunes']} ({mejor_coincidencia['por_nombre_cientifico']})"
                            nombre_cientifico = mejor_coincidencia["por_nombre_cientifico"]
                        else:
                            # Si no hay coincidencia suficiente, usar el tipo detectado pero dejar la especie sin identificar
                            tipo_seleccionado = tipo_detectado
                            especie_seleccionada = "Especie no identificada"
                            nombre_cientifico = "Nombre cient√≠fico no disponible"

                    except Exception as e:
                        st.error(f"Error al analizar la imagen con Gemini AI: {e}")
                        tipo_seleccionado = "Desconocido"
                        especie_seleccionada = "Especie no identificada"
                        nombre_cientifico = "Nombre cient√≠fico no disponible"

                # Asignar valores a session_state
                st.session_state[f"tipo_{i}"] = tipo_seleccionado
                st.session_state[f"especie_{i}"] = especie_seleccionada

                # Mostrar los valores asignados autom√°ticamente
                st.write(f"**Tipo:** {tipo_seleccionado}" if tipo_seleccionado != "Desconocido" else "")
                st.write(f"**Especie:** {especie_seleccionada}" if especie_seleccionada != "Especie no identificada" else "")

                if tipo_seleccionado == "Desconocido":
                    # Permitir selecci√≥n manual del Tipo y Especie
                    tipo_seleccionado_manual = st.selectbox(
                        f"Selecciona el tipo para la imagen {i+1}:",
                        tipos_especies,
                        key=f"tipo_manual_{i}",
                        index=tipos_especies.index("PESCADOS") if "PESCADOS" in tipos_especies else 0
                    )
                    st.session_state[f"tipo_{i}"] = tipo_seleccionado_manual

                    # Filtrar especies por el tipo seleccionado manualmente
                    especies_filtradas_manual = [
                        especie for especie in especies
                        if especie["por_tipo"] == tipo_seleccionado_manual
                    ]

                    # Crear lista de opciones para el selectbox
                    opciones_especies_manual = [
                        f"{especie['por_nombres_comunes']} ({especie['por_nombre_cientifico']})"
                        for especie in especies_filtradas_manual
                    ]

                    # Selectbox para elegir la especie manualmente
                    especie_seleccionada_manual = st.selectbox(
                        f"Selecciona la especie para la imagen {i+1}:",
                        opciones_especies_manual,
                        key=f"especie_manual_{i}"
                    )
                    st.session_state[f"especie_{i}"] = especie_seleccionada_manual

                    # Mostrar el nombre cient√≠fico
                    nombre_cientifico_manual = next(
                        (especie["por_nombre_cientifico"] for especie in especies_filtradas_manual
                         if f"{especie['por_nombres_comunes']} ({especie['por_nombre_cientifico']})" == especie_seleccionada_manual),
                        "Nombre cient√≠fico no disponible"
                    )
                    st.write(f"**Nombre cient√≠fico:** {nombre_cientifico_manual}")
                else:
                    # Si el tipo no es desconocido, permitir al usuario cambiar la especie si lo desea
                    st.write("**Nota:** La especie ha sido asignada autom√°ticamente. Si necesitas cambiarla, selecci√≥nala manualmente a continuaci√≥n.")
                    
                    # Filtrar especies por el tipo asignado autom√°ticamente
                    especies_filtradas_auto = [
                        especie for especie in especies
                        if especie["por_tipo"] == tipo_seleccionado
                    ]

                    # Crear lista de opciones para el selectbox
                    opciones_especies_auto = [
                        f"{especie['por_nombres_comunes']} ({especie['por_nombre_cientifico']})"
                        for especie in especies_filtradas_auto
                    ]

                    # Selectbox para elegir la especie manualmente (opcional)
                    especie_seleccionada_manual_auto = st.selectbox(
                        f"Si deseas cambiar la especie para la imagen {i+1}, selecci√≥nala aqu√≠:",
                        opciones_especies_auto,
                        key=f"especie_manual_auto_{i}",
                        index=opciones_especies_auto.index(especie_seleccionada) if especie_seleccionada in opciones_especies_auto else 0
                    )
                    st.session_state[f"especie_{i}"] = especie_seleccionada_manual_auto

                    # Mostrar el nombre cient√≠fico
                    nombre_cientifico_manual_auto = next(
                        (especie["por_nombre_cientifico"] for especie in especies_filtradas_auto
                         if f"{especie['por_nombres_comunes']} ({especie['por_nombre_cientifico']})" == especie_seleccionada_manual_auto),
                        "Nombre cient√≠fico no disponible"
                    )
                    st.write(f"**Nombre cient√≠fico:** {nombre_cientifico_manual_auto}")

        # Despu√©s de procesar todas las im√°genes, mostrar el logo
        st.markdown("### Powered by Gemini")

        # Ruta al archivo JSON de Lottie (aseg√∫rate de que el archivo exista en esta ubicaci√≥n)
        lottie_path = os.path.join(base_dir, "gemini_logo.json")
        gemini_logo = load_lottie_file(lottie_path)

        # Mostrar la animaci√≥n Lottie con tama√±o peque√±o
        st_lottie.st_lottie(
            gemini_logo, 
            key='logo', 
            height=50,  # Ajusta la altura seg√∫n tus necesidades
            width=50,   # Ajusta el ancho seg√∫n tus necesidades
            loop=True,
            quality="low"  # Opcional: reduce la calidad para mejorar el rendimiento
        )
 
    # Formulario para datos del entregable
    st.header("Datos del entregable")
    #firma = seccion_firma()
    # Agregar nuevos campos al formulario
    dni = st.text_input("Introduce tu DNI", max_chars=8)
    # Variables para almacenar datos obtenidos de la API
    nombres = ""
    iniciales = ""
    ruc = ""
    # Bot√≥n para obtener datos de SUNAT
    if dni:
        if len(dni) == 8:
            nombres, iniciales, ruc = obtener_datos_sunat(dni)
            if nombres and iniciales and ruc:
                st.write(f"**Nombre Completo:** {nombres}")
                st.write(f"**Iniciales:** {iniciales}")
                st.write(f"**RUC:** {ruc}")
        else:
            st.error("El DNI debe tener exactamente 8 d√≠gitos.")
    os_input = st.text_input("Introduce tu OS (4 d√≠gitos)", max_chars=4)        
    email = st.text_input("Introduce tu Email", max_chars=50)
    telefono = st.text_input("Tel√©fono", max_chars=15)
    seccion_direccion()
    # Validaciones b√°sicas
    if dni and not dni.isdigit():
        st.error("El DNI debe contener solo n√∫meros.")
    if telefono and not telefono.isdigit():
        st.error("El tel√©fono debe contener solo n√∫meros.")
    if os_input and not os_input.isdigit():
        st.error("La OS debe contener solo n√∫meros.")
    if os_input and len(os_input) > 4:
        st.error("La OS debe tener como m√°ximo 4 d√≠gitos.")
    # Lista desplegable de tipos de entregable
    tipo_entregable = st.selectbox(
        "Selecciona el tipo de entregable",
        ["Unico", "Primer", "Segundo", "Tercer", "Cuarto"],
    )
    # Asignar valores a variables
    if tipo_entregable in ["Unico", "Primer"]:
        n = 1
    elif tipo_entregable == "Segundo":
        n = 2
    elif tipo_entregable == "Tercer":
        n = 3
    else:
        n = 4
    # Reemplazo de selecci√≥n de mes y d√≠a por un calendario
    fecha_seleccionada = st.date_input(
        "Selecciona la fecha de presentaci√≥n",
        value=datetime.now(),
        min_value=datetime(2000, 1, 1),
        max_value=datetime(2100, 12, 31)
    )

    # Extraer mes, d√≠a y a√±o de la fecha seleccionada
    mes_num = fecha_seleccionada.month
    mes_seleccionado = MESES_ES.get(mes_num, fecha_seleccionada.strftime("%B"))
    dia_seleccionado = fecha_seleccionada.day
    year = fecha_seleccionada.year

    # Implementaci√≥n en el formulario
    st.subheader("Selecciona la Ciudad y Departamento en el que desempe√±as tus funciones")
    # Cargar el archivo ciudades.json
    ciudades = cargar_ciudades()
    departamentos = list(ciudades.keys()) + ["Ingreso manual"]
    departamento = st.selectbox("Selecciona el Departamento:", departamentos)
    # Mostrar ciudades o campo editable
    if departamento == "Ingreso manual":
        ciudad = st.text_input("Escribe el nombre de tu ciudad:")
    else:
        ciudad = st.selectbox("Selecciona la Ciudad:", ciudades[departamento])
    st.write(f"Departamento seleccionado: {departamento}, Ciudad seleccionada: {ciudad}")
    # L√≥gica para seleccionar el banco y autorrellenar el CCI en funci√≥n del banco y cuenta
    st.subheader("Datos Bancarios")
    st.write("_Los datos bancarios son opcionales, solo se utilizan para generar el_ Formato 1 DJ Menores a 8 UIT")
    st.write("_De no ingresar sus datos bancarios estos no seran reemplazados en_ Formato 1 DJ Menores a 8 UIT")
    banco_seleccionado = st.selectbox("Selecciona tu banco:", ["BCP", "Interbank", "Scotiabank", "Banco de la Naci√≥n", "BanBif", "Otros"])
    cuenta = st.text_input("Ingresa tu cuenta", max_chars=20, placeholder="Ej. 123456789 o 123-456-78")

    # Generar el CCI basado en el banco y cuenta ingresada, si corresponde
    cci = ""
    if banco_seleccionado != "Otros" and cuenta:
        cuenta_limpia = cuenta.replace("-", "")
        if banco_seleccionado == "BCP":
            cci = "002" + cuenta_limpia + "13"
        elif banco_seleccionado == "Interbank":
            cci = "003" + cuenta_limpia + "43"
        elif banco_seleccionado == "Scotiabank":
            cci = "00936020" + cuenta_limpia + "95"
        elif banco_seleccionado == "Banco de la Naci√≥n":
            cci = "018-781-0" + cuenta_limpia + "-55"
        elif banco_seleccionado == "BanBif":
            cci = "0386501" + cuenta_limpia + "83"

    # Mostrar el CCI en un campo editable
    cci = st.text_input("CCI (editable):", value=cci)
    st.write(f"Banco seleccionado: {banco_seleccionado}, CCI generado: {cci}")

    # Subheader para las recomendaciones
    st.subheader("Ingrese su recomendaci√≥n al Informe")

    # Lista desplegable para seleccionar el tipo de recomendaci√≥n
    opcion_recomendacion = st.selectbox(
        "Selecciona el tipo de recomendaci√≥n:",
        ["Generar recomendaciones por IA", "Ingresar recomendaciones manualmente"]
    )

    if opcion_recomendacion == "Generar recomendaciones por IA":
        # Bot√≥n para generar recomendaciones
        if st.button("Generar recomendaciones por IA"):
            # Validaciones previas
            if not excel_files:
                st.error("Por favor, sube al menos un archivo Excel antes de generar recomendaciones.")
            elif not pdf_file:
                st.error("Por favor, sube el archivo PDF de TDR antes de generar recomendaciones.")
            else:
                try:
                    # Crear directorios temporales para almacenar archivos
                    with tempfile.TemporaryDirectory() as carpeta_data:
                        # Guardar los archivos Excel subidos en la carpeta de datos
                        for file in excel_files:
                            with open(os.path.join(carpeta_data, file.name), "wb") as f:
                                f.write(file.getbuffer())
                        st.success("Archivos Excel guardados correctamente.")

                        # Procesar los archivos Excel para generar data.xlsx
                        output_excel = os.path.join(carpeta_data, "data.xlsx")
                        success = generar_data_excel(
                            [os.path.join(carpeta_data, file.name) for file in excel_files],
                            output_excel,
                        )
                        if success:
                            st.success("Archivo data.xlsx generado correctamente.")
                            # Leer data.xlsx y almacenarlo en session_state
                            with open(output_excel, "rb") as f:
                                st.session_state['data_excel'] = f.read()
                            st.success("Archivo data.xlsx almacenado en la sesi√≥n correctamente.")
                        else:
                            st.error("No se pudo generar data.xlsx.")

                        # Asegurarse de que data_excel no es None y es de tipo bytes
                        if st.session_state['data_excel'] and isinstance(st.session_state['data_excel'], bytes):
                            with open(output_excel, "rb") as f:
                                st.session_state['data_excel'] = f.read()
                                logging.info(f"Tipo de st.session_state['data_excel']: {type(st.session_state['data_excel'])}")
                        else:
                            st.error("El archivo data.xlsx no se gener√≥ correctamente o no est√° en el formato esperado.")

                        # Guardar data.xlsx en session_state como BytesIO
                        with open(output_excel, "rb") as f:
                            st.session_state['data_excel'] = f.read()
                            logging.info(f"Tipo de st.session_state['data_excel']: {type(st.session_state['data_excel'])}")

                        # **Guardar el PDF subido en un archivo temporal y definir tmp_pdf_path**
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
                            tmp_pdf.write(pdf_file.getbuffer())
                            tmp_pdf_path = tmp_pdf.name
                        st.success("PDF de TDR guardado correctamente.")

                        # Abrir el archivo temporal en modo binario
                        with open(tmp_pdf_path, "rb") as f_pdf:
                            servicio_extraido = extraer_nombre_servicio(f_pdf)
                            f_pdf.seek(0)  # Reiniciar el cursor para la siguiente lectura
                            actividades_extraidas = extraer_actividades(f_pdf)
                            if servicio_extraido and actividades_extraidas:
                                st.session_state['servicio'] = servicio_extraido
                                st.session_state['actividades'] = {
                                    "a": actividades_extraidas[0],
                                    "b": actividades_extraidas[1],
                                    "c": actividades_extraidas[2],
                                    "d": actividades_extraidas[3],
                                    "e": actividades_extraidas[4],
                                }
                                logging.info("Datos de servicio y actividades extra√≠dos correctamente.")
                            else:
                                st.error("Error al extraer datos de servicio y actividades del PDF.")

                        # **Eliminar el archivo temporal despu√©s de usarlo (opcional)**
                        #os.remove(tmp_pdf_path)

                        # Guardar 'servicio' y actividades en session_state
                        st.session_state['servicio'] = servicio_extraido
                        st.session_state['actividades'] = {
                            "a": actividades_extraidas[0],
                            "b": actividades_extraidas[1],
                            "c": actividades_extraidas[2],
                            "d": actividades_extraidas[3],
                            "e": actividades_extraidas[4],
                        }
                        # Despu√©s de extraer los datos
                        if servicio_extraido and all(isinstance(v, str) for v in actividades_extraidas):
                            st.session_state['servicio'] = servicio_extraido
                            st.session_state['actividades'] = {
                                "a": actividades_extraidas[0],
                                "b": actividades_extraidas[1],
                                "c": actividades_extraidas[2],
                                "d": actividades_extraidas[3],
                                "e": actividades_extraidas[4],
                            }
                            logging.info(f"Servicio guardado: {servicio_extraido}")
                            logging.info(f"Actividades guardadas: {st.session_state['actividades']}")

                        # Leer y formatear los dataframes
                        dataframes_text = leer_y_formatear_dataframes(BytesIO(st.session_state['data_excel']))

                        # Crear el prompt para la IA
                        prompt = f"""
                        Bas√°ndote en los siguientes datos, genera recomendaciones para el informe sin repetir la informaci√≥n proporcionada.

                        Datos Principales:
                        Servicio: {st.session_state['servicio']}
                        Actividades:
                        A: {st.session_state['actividades']['a']}
                        B: {st.session_state['actividades']['b']}
                        C: {st.session_state['actividades']['c']}
                        D: {st.session_state['actividades']['d']}
                        E: {st.session_state['actividades']['e']}

                        Datos Adicionales:
                        {dataframes_text}

                        Por favor, elabora recomendaciones basadas en esta informaci√≥n.
                        """

                        # Generar la respuesta de la IA
                        recomendaciones_ia = generar_respuesta_rapidapi(prompt)

                        # Almacenar las recomendaciones en session_state
                        st.session_state['recomendaciones'] = recomendaciones_ia

                        st.success("Recomendaciones generadas exitosamente.")

                except Exception as e:
                    st.error(f"Error al generar recomendaciones por IA: {e}")

        # Mostrar el campo de texto editable con las recomendaciones generadas por la IA
        recomendaciones = st.text_area(
            "Recomendaciones Generadas por IA (puedes editar):",
            value=st.session_state.get('recomendaciones', ''),
            height=200,
            help="Puedes usar formato markdown: **negrita**, *cursiva*, ‚Ä¢ lista, etc."
        )
        # Agregar esta l√≠nea para mostrar la vista previa
        st.markdown("### Vista previa del formato:")
        st.markdown(recomendaciones)

        # Actualizar session_state con cualquier edici√≥n del usuario
        st.session_state['recomendaciones'] = recomendaciones

    else:
        # Campo editable con la recomendaci√≥n manual prellenada
        recomendaciones_manual = (
            "Coordinar con Capitan√≠a de Puerto, respecto al caso de las matr√≠culas y permisos "
            "de pesca de las embarcaciones pesqueras artesanales que operan en la zona, para que "
            "se formalicen pues hay un porcentaje de las mismas que a√∫n no cuentan con estos documentos, "
            "pese a las amnist√≠as que se les ha otorgado."
        )
        recomendaciones = st.text_area(
            "Ingrese sus recomendaciones manualmente:",
            value=recomendaciones_manual,
            height=200
        )
        # Almacenar las recomendaciones manuales en session_state
        st.session_state['recomendaciones'] = recomendaciones

    # Variables para almacenar datos extra√≠dos del PDF
    servicio = ""
    act_a = act_b = act_c = act_d = act_e = ""

    # Bot√≥n para generar el entregable
    if st.button("Generar entregable"):
        # Validaciones
        if not excel_files:
            st.error("Por favor, sube al menos un archivo Excel para continuar.")
        elif not (telefono and dni and email and os_input and st.session_state['direccion']):
            st.error("Por favor, completa todos los campos del formulario.")
        elif not ruc:
            st.error("No se pudo obtener el RUC. Verifica el DNI ingresado.")
        elif not pdf_file:
            st.error("Por favor, sube el archivo PDF de TDR.")
        elif not uploaded_images:
            st.error("Por favor, sube al menos una imagen.")
        elif not st.session_state['recomendaciones'].strip():
            st.error("Las recomendaciones no pueden estar vac√≠as. Por favor, genera o ingresa recomendaciones.")
        else:
            try:
                # Crear directorios temporales para almacenar archivos
                with tempfile.TemporaryDirectory() as carpeta_data, tempfile.TemporaryDirectory() as carpeta_recursos:
                    # Guardar los archivos Excel subidos en la carpeta de datos
                    for file in excel_files:
                        with open(os.path.join(carpeta_data, file.name), "wb") as f:
                            f.write(file.getbuffer())
                    st.success("Archivos Excel guardados correctamente.")

                    # Verificar si data.xlsx ya fue generado al generar recomendaciones
                    if st.session_state['data_excel']:
                        logging.info(f"Tipo de st.session_state['data_excel']: {type(st.session_state['data_excel'])}")
                        # Guardar data.xlsx desde session_state
                        output_excel = os.path.join(carpeta_data, "data.xlsx")
                        with open(output_excel, "wb") as f:
                            f.write(st.session_state['data_excel'])
                        st.success("Archivo data.xlsx recuperado desde recomendaciones generadas.")
                    else:
                        # Generar data.xlsx si no existe
                        output_excel = os.path.join(carpeta_data, "data.xlsx")
                        success = generar_data_excel(
                            [os.path.join(carpeta_data, file.name) for file in excel_files],
                            output_excel,
                        )
                        if success:
                            st.success("Archivo data.xlsx generado correctamente.")
                            with open(output_excel, "rb") as f:
                                st.session_state['data_excel'] = f.read()
                            st.success("Archivo data.xlsx almacenado en la sesi√≥n correctamente.")
                        else:
                            st.error("No se pudo generar data.xlsx.")

                    # Copiar data.xlsx a carpeta_recursos
                    shutil.copy(output_excel, os.path.join(carpeta_recursos, "data.xlsx"))
                    st.success("Archivo data.xlsx copiado correctamente a la carpeta de recursos.")

                    # Guardar las im√°genes subidas
                    for image in uploaded_images:
                        with open(os.path.join(carpeta_recursos, image.name), "wb") as f:
                            f.write(image.getbuffer())
                    st.success("Im√°genes guardadas correctamente.")

                    # Guardar el PDF subido en un archivo temporal
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
                        tmp_pdf.write(pdf_file.getbuffer())
                        pdf_path = tmp_pdf.name
                    st.success("PDF de TDR guardado correctamente.")

                    # Extraer datos del PDF
                    st.header("Extrayendo datos del PDF")
                    # Abrir el archivo temporal en modo binario
                    with open(pdf_path, "rb") as f_pdf:
                        servicio_extraido = extraer_nombre_servicio(f_pdf)
                        f_pdf.seek(0)  # Reiniciar el cursor para la siguiente lectura
                        actividades_extraidas = extraer_actividades(f_pdf)

                    # Contar el n√∫mero de reportes Excel
                    reportes = contar_archivos(excel_files)
                    st.write(f"**N√∫mero de Reportes Excel:** {reportes}")

                    # Formatear la OS para que tenga 4 d√≠gitos
                    orden = os_input.zfill(4)
                    st.write(f"**Orden de Servicio (OS):** {orden}")

                    # Generar el gr√°fico de barras
                    try:
                        df_grafico = pd.read_excel(output_excel, sheet_name="grafico")
                        grafico_path = generar_grafico(df_grafico, carpeta_recursos)
                        st.image(
                            grafico_path,
                            caption="Gr√°fico de Barras Generado",
                            use_column_width=True,
                        )
                        st.success("Gr√°fico generado correctamente.")
                    except Exception as e:
                        st.error(f"No se pudo generar el gr√°fico: {e}")
                        grafico_path = None

                    # Definir los valores para reemplazar en el documento
                    valores = {
                        "nombres": nombres,
                        "iniciales": iniciales,
                        "ruc": ruc,
                        "telefono": telefono,
                        "email": email,
                        "dni": dni,
                        "direccion": st.session_state.get('direccion', ''),
                        "dia": dia_seleccionado,
                        "mes": mes_seleccionado,
                        "year": year,
                        "n": n,
                        "entregable": tipo_entregable,
                        "orden": orden,
                        "reportes": reportes,
                        "servicio": st.session_state.get('servicio', ''),
                        "act_a": st.session_state.get('actividades', {}).get('a', ''),
                        "act_b": st.session_state.get('actividades', {}).get('b', ''),
                        "act_c": st.session_state.get('actividades', {}).get('c', ''),
                        "act_d": st.session_state.get('actividades', {}).get('d', ''),
                        "act_e": st.session_state.get('actividades', {}).get('e', ''),
                        "ciudad": ciudad,
                        "banco": banco_seleccionado,
                        "cuenta": cuenta,
                        "cci": cci,
                        "recomendaciones": markdown_to_docx_text(st.session_state.get('recomendaciones', '')),  # Incluir recomendaciones
                        #"firma": firma,
                    }
                    
                    # Leer y formatear los dataframes
                    if st.session_state['data_excel']:
                        excel_bytes_io = BytesIO(st.session_state['data_excel'])
                        descripcion_df = pd.read_excel(excel_bytes_io, sheet_name="descripcion", engine='openpyxl')
                        total_df = pd.read_excel(excel_bytes_io, sheet_name="total", engine='openpyxl')
                        especies_df = pd.read_excel(excel_bytes_io, sheet_name="especies", engine='openpyxl')
                        procedencia_df = pd.read_excel(excel_bytes_io, sheet_name="procedencia", engine='openpyxl')
                        comentarios_df = pd.read_excel(excel_bytes_io, sheet_name="comentarios", engine='openpyxl')
                    else:
                        st.error("El archivo data.xlsx no est√° disponible en la sesi√≥n.")
                        return

                    # Verificar contenido de los DataFrames
                    logging.info("Contenido de 'descripcion_df':")
                    logging.info(descripcion_df.head())
                    st.write("DESARROLLO DE ACTIVIDADES:", descripcion_df[['Fecha', 'Hora', 'Descripcion']].head())
                    logging.info("Contenido de 'total_df':")
                    logging.info(total_df.head())
                    st.write("PRINCIPALES ESPECIES DESEMBARCADAS:", total_df.head())
                    logging.info("Contenido de 'especies_df':")
                    logging.info(especies_df.head())
                    st.write("VARIACION DE PRECIO DE LAS PRINCIPALES ESPECIES DESCARGAS:", especies_df[['NOMBRE COMUN', 'NOMBRE CIENTIFICO', 'PRECIO MAS BAJO (S/)', 'PRECIO MAS ALTO (S/)']].head())
                    logging.info("Contenido de 'procedencia_df':")
                    logging.info(procedencia_df.head())
                    st.write("Las zonas de pesca fueron:", procedencia_df.head())
                    # Obtener los comentarios de la hoja 'comentarios'
                    comentarios = obtener_comentarios_excel(comentarios_df)

                    # Actualizar los valores para reemplazar
                    valores.update({
                        'servicio': servicio_extraido,  # Usar el valor extra√≠do
                        'act_a': actividades_extraidas[0],
                        'act_b': actividades_extraidas[1],
                        'act_c': actividades_extraidas[2],
                        'act_d': actividades_extraidas[3],
                        'act_e': actividades_extraidas[4],
                        'Especie_principal_1': comentarios.get('Especie_principal_1', ''),
                        'Especie_principal_2': comentarios.get('Especie_principal_2', ''),
                        'Especie_principal_3': comentarios.get('Especie_principal_3', ''), 
                        'Especie_minima_1': comentarios.get('Especie_minima_1', ''),
                        'Especie_minima_2': comentarios.get('Especie_minima_2', ''),
                        'Especie_minima_3': comentarios.get('Especie_minima_3', ''),         
                        'Mayor_dia': comentarios.get('Mayor_dia', ''),
                        'Menor_dia': comentarios.get('Menor_dia', ''),
                        'Aparejos': comentarios.get('Aparejos', ''),
                        'Mayor_Precio_1': comentarios.get('Mayor_Precio_1', ''),
                        'Mayor_Precio_2': comentarios.get('Mayor_Precio_2', ''),
                        'Mayor_Precio_3': comentarios.get('Mayor_Precio_3', ''),
                        'Menor_Precio_1': comentarios.get('Menor_Precio_1', ''),
                        'Menor_Precio_2': comentarios.get('Menor_Precio_2', ''),
                        'Menor_Precio_3': comentarios.get('Menor_Precio_3', ''),
                        'Total': comentarios.get('Total', '')  
                    })

                    # Definir el nombre del archivo de salida
                    archivo_salida = f"7. Informe N¬∞ 00{n}-{year}-{iniciales}.docx"
                    archivo_entrada = os.path.join(base_dir, "Informe.docx")  # Plantilla en la ra√≠z
                    ruta_salida = os.path.join(carpeta_recursos, archivo_salida)

                    # Verificar si la plantilla de Word existe
                    if not os.path.exists(archivo_entrada):
                        st.error(f"La plantilla {archivo_entrada} no se encontr√≥.")
                    else:
                        # Llenar la plantilla de Word usando docxtpl
                        llenar_plantilla_word(
                            archivo_entrada, ruta_salida, valores, grafico_path
                        )

                        # Llenar la plantilla y las tablas
                        llenar_plantilla_word(
                            archivo_entrada=archivo_entrada,
                            archivo_salida=ruta_salida,
                            valores=valores,  # Aseg√∫rate de que 'valores' contiene todas las claves necesarias
                            grafico_path=grafico_path,
                            df_descripcion=descripcion_df,
                            df_total=total_df,
                            df_especies=especies_df,
                            df_procedencia=procedencia_df,
                            uploaded_images=uploaded_images,
                            imagenes_dir=carpeta_recursos
                        )

                        # Convertir el documento de Word a PDF
                        archivo_pdf_salida = archivo_salida.replace(".docx", ".pdf")
                        ruta_pdf_salida = os.path.join(carpeta_recursos, archivo_pdf_salida)
                        convertir_a_pdf(ruta_salida, ruta_pdf_salida)

                        # Procesar y convertir documentos adicionales a PDF
                        documentos_adicionales = [
                            "4. Formato 1 DJ Menores a 8 UIT.docx",
                            "Cargo.docx",
                            "Carta de Presentacion.docx",
                        ]
                        procesar_documentos_adicionales(
                            documentos_adicionales, carpeta_recursos, valores, base_dir
                        )

                        # Empaquetar todos los archivos generados en un ZIP
                        archivos_para_zip = [
                            archivo_salida, 
                            archivo_pdf_salida, 
                            "data.xlsx"  # A√±ade data.xlsx aqu√≠
                        ] + [
                            docx.replace(".docx", ".pdf") for docx in documentos_adicionales
                        ]
                        zip_buffer = crear_zip_archivos(archivos_para_zip, carpeta_recursos)

                        # Proveer enlaces de descarga
                        st.header("Descarga de Archivos Generados")
                        st.download_button(
                            label="Descargar Todos los Archivos Generados (ZIP)",
                            data=zip_buffer,
                            file_name="archivos_generados.zip",
                            mime="application/zip",
                        )

                        st.success("Entregable generado exitosamente.")

            except Exception as e:
                st.error(f"Ocurri√≥ un error durante el procesamiento: {e}")
                st.exception(e)

    crear_donation_footer(base_dir)

if __name__ == "__main__":
    main()
